from docx.document import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns

def add_page_number(run) -> None:
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def set_document(doc, title, heading) -> None:
    par = doc.add_paragraph(title)
    par.style = doc.styles['Title']
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = heading
    add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
    doc.add_section(WD_SECTION.CONTINUOUS)
    section = doc.sections[1]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')