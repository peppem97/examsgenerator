from docx import Document
from utility import *
import random

class Exam():
    def __init__(self, config, all_questions) -> None:
        self.config = config
        self.questions = all_questions
        self.doc = Document()
        self.doc_solutions = Document()
        set_document(self.doc, self.config['title'], self.config['heading'])
        set_document(self.doc_solutions, self.config['title'], self.config['heading'])


    def shuffle_questions(self):
        if self.config['shuffle_questions']:
            random.shuffle(self.questions)
        if self.config['shuffle_options']:
            for question in self.questions:
                random.shuffle(question['options'])
        self.questions = self.questions[0: self.config['questions_number']]


    def write_exam(self, exam_number) -> None:
        for index, question in enumerate(self.questions):
            self.doc.add_heading(f"{index + 1}) {question['question']}", 3)
            if question['type'] == "QUIZ":
                for i in range(0, self.config['options_supported']):
                    self.doc.add_paragraph(style='List Bullet').add_run(question['options'][i]['text'])
        self.doc.save(f"{self.config['destination_file']}_{str(exam_number + 1)}.docx")

    
    def write_exam_with_solutions(self, exam_number) -> None:
        for index, question in enumerate(self.questions):
            self.doc_solutions.add_heading(f"{index + 1}) {question['question']}", 3)
            if question['type'] == "QUIZ":
                for i in range(0, self.config['options_supported']):
                    runner = self.doc_solutions.add_paragraph(style='List Bullet').add_run(question['options'][i]['text'])
                    runner.bold = question['options'][i]['correct']
        self.doc_solutions.save(f"{self.config['destination_file']}_{str(exam_number + 1)}_solutions.docx")


