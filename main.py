import pandas as pd
import random
import json
from docx import Document
from docx.document import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns


class ExamsGenerator():
    def __init__(self) -> None:
        with open('config.json', 'r') as f:
            self.config = json.load(f)
        self.df = pd.read_excel(f"{self.config['source_file']}.{self.config['source_extension']}")
        self.show_params()


    def init_document(self):
        self.doc = Document()
        if self.config['export_solutions']:
            self.doc_solutions = Document()


    def show_params(self):
        print("********************")
        print(f"Materia: {self.config['subject']}")
        print(f"Classe: {self.config['classroom']}")
        print("********************")
            
                
    def add_page_number(self, run) -> None:
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


    def set_document(self, doc) -> None:
        par = doc.add_paragraph(self.config['title'])
        par.style = doc.styles['Title']
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = self.config['heading']
        self.add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
        doc.add_section(WD_SECTION.CONTINUOUS)
        section = doc.sections[1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')

    
    def check_row(self, row) -> None:
        if self.config['deep_filtering']:
            return (
                row[self.config['subject_denomination']] == self.config['subject'] and 
                row[self.config['classroom_denomination']] == self.config['classroom'] and 
                row[self.config['era_denomination']] == self.config['era'] and
                row[self.config['sector_denomination']] == self.config['sector'])
        else:
            return True


    def pool_questions(self) -> None:
        self.questions = []
        for _, row in self.df.iterrows():
            if (self.check_row(row)):
                question = {
                        "question": str(row[self.config['question_denomination']]), 
                        "options": []
                    }
                for i in range(0, self.config['options_supported']):
                    question["options"].append({"text": str(row[f'{self.config["option_denomination"]}_{i + 1}']), "correct": True if int(row[self.config['solution_denomination']]) == (i + 1) else False})
                self.questions.append(question)


    def count_questions(self) -> None:
        print(f"Numero di domande: {len(self.questions)}")
        print("********************")


    def filter_questions(self):
        if self.config['shuffle_questions']:
            random.shuffle(self.questions)
        if self.config['shuffle_options']:
            for question in self.questions:
                random.shuffle(question['options'])
        self.questions = self.questions[0: self.config['questions_number']]


    def write_exam(self, number_exam) -> None:
        for index, question in enumerate(self.questions):
            self.doc.add_heading(f"{index + 1}) {question['question']}", 3)
            for i in range(0, self.config['options_supported']):
                self.doc.add_paragraph(style='List Bullet').add_run(question['options'][i]['text'])
        self.doc.save(f"{self.config['destination_file']}_{str(number_exam + 1)}.docx")

    
    def write_exam_with_solutions(self, number_exam) -> None:
        for index, question in enumerate(self.questions):
            self.doc_solutions.add_heading(f"{index + 1}) {question['question']}", 3)
            for i in range(0, self.config['options_supported']):
                runner = self.doc_solutions.add_paragraph(style='List Bullet').add_run(question['options'][i]['text'])
                runner.bold = question['options'][i]['correct']
        self.doc_solutions.save(f"{self.config['destination_file']}_{str(number_exam + 1)}_solutions.docx")

            
    def start(self) -> None:
        for i in range (0, self.config["exams_number"]):
            self.init_document()
            self.set_document(self.doc)
            if self.config['export_solutions']:
                self.set_document(self.doc_solutions)
            self.pool_questions()
            if self.config['only_count']:
                self.count_questions()
            else:
                self.filter_questions()
                self.write_exam(i)
                if self.config['export_solutions']:
                    self.write_exam_with_solutions(i)
                print("Esami generati!")
                print("********************")


if __name__ == "__main__":
    exams_generator = ExamsGenerator()
    exams_generator.start()


