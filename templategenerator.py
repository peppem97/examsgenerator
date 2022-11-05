from docx import Document
from docx.document import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class TemplateGenerator():
    def __init__(self, title) -> None:
        self.doc = Document()
        par = self.doc.add_paragraph(f"{title}")
        par.style = self.doc.styles['Title']
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section = self.doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Cognome e Nome: ___________________________________________"
        self.doc.add_section(WD_SECTION.CONTINUOUS)
        section = self.doc.sections[1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')

    def write_template(self, name):
        self.doc.save(f"{name}.docx")